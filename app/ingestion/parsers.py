"""Document parsers for PDF, Word, and spreadsheet files.

These are provided as working boilerplate so the candidate can focus on the
agent layer. They return plain text (and light structure) that the chunker
turns into indexable units. Candidates may extend these (e.g. table-aware
extraction, OCR) but are not required to.
"""
from pathlib import Path

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.schemas import DocumentType


def detect_type(path: Path) -> DocumentType:
    suffix = path.suffix.lower()
    return {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOCX,
        ".xlsx": DocumentType.XLSX,
        ".xls": DocumentType.XLSX,
        ".csv": DocumentType.XLSX,
    }.get(suffix, DocumentType.UNKNOWN)


def parse_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        pages.append(f"[page {i + 1}]\n{text}")
    return "\n\n".join(pages)


def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table cell text as well.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_spreadsheet(path: Path) -> str:
    """Flatten every sheet into a text representation the model can read."""
    if path.suffix.lower() == ".csv":
        frames = {"Sheet1": pd.read_csv(path)}
    else:
        frames = pd.read_excel(path, sheet_name=None)  # dict[str, DataFrame]

    parts = []
    for sheet_name, df in frames.items():
        parts.append(f"[sheet: {sheet_name}]")
        parts.append(df.to_csv(index=False))
    return "\n\n".join(parts)


def parse(path: Path) -> tuple[str, DocumentType]:
    doc_type = detect_type(path)
    if doc_type == DocumentType.PDF:
        return parse_pdf(path), doc_type
    if doc_type == DocumentType.DOCX:
        return parse_docx(path), doc_type
    if doc_type == DocumentType.XLSX:
        return parse_spreadsheet(path), doc_type
    raise ValueError(f"Unsupported file type: {path.suffix}")
